import streamlit as st
import random
import numpy as np
from langchain.chains.question_answering import load_qa_chain
from langchain.llms import OpenAI
from langchain.chains import RetrievalQA
from langchain.callbacks import get_openai_callback
from langchain.vectorstores import Chroma
from docx import Document
import io
from utils import *
from prompt import *

def main():

    #Create Session State Variable
    if 'OpenAI_API_Key' not in st.session_state:
        st.session_state['OpenAI_API_Key'] = ''

    st.set_page_config(page_title="Multiple Choice Quiz Generator")
    st.title("The AI Quiz Creator")
    st.subheader("Create Quizzes with the help of AI")
    
    st.markdown("""
                <style>
                div.stButton > button:first-child {
                    background-color: #0000ff;
                    color: #ffffff;
                }
                </style>
                """, unsafe_allow_html=True)
    
    # upload informational pdf files
    pdf_docs = st.file_uploader("Upload the pdf for your quiz here", type=["pdf"], accept_multiple_files=True)

    # quiz_description = st.text_area("Paste your learning goals here (i.e. benchmarks or standards) or just say 'give me some questions!'", key = "1")

    st.sidebar.title("Pass the Keys, Please")
    st.session_state["OpenAI_API_Key"] = st.sidebar.text_input("What's your OpenAI API Key?", type = "password")

    options = st.selectbox(
        "How would you like your questions?",
        ("Key Terms", "Keyterm Finder", "Standard/Benchmark", "Random questions", "Sequential questions")
    )

    if st.session_state["OpenAI_API_Key"] != "" and options == "Sequential questions":

        if pdf_docs != None:
            # with st.spinner("Wait for it..."):

            #Extract the text from the uploaded pdf file
            text = pdf_to_text(pdf_docs)

            #Split into chunks
            chunks = split_text(text, 750, 50)

            question_count = st.slider("Number of Questions", 1, len(chunks))

            st.write("Proportions must add up to 1")
            long = st.text_input("Proportion of scenario questions")
            short = st.text_input("Proportion of short questions")
            blank = st.text_input("Proportion of fill in blank questions")


            submit = st.button("Generate Quiz!")

            if submit:
                with get_openai_callback() as cb:
                    question_num = 1
                    text_to_download = Document()
                    for i in range(question_count):
                        try:
                            # question_type = random.randint(0,2)
                            question_type = np.random.choice(np.arange(0,3), p=[float(long), float(short), float(blank)])
                            # st.write(question_type)
                            
                            format_instructions = formats[question_type]
                            
                            if question_type == 0:

                                scenario, questions, choices, answers, explanations = prompts[question_type](format_instructions, chunks[i])
                                    
                                st.write(str(question_num) + ": " + scenario + " " + questions)
                                st.write(choices)
                                st.write("")
                                # st.write(answer)
                                with st.expander('Show me the answer 👀'): 
                                    st.write("Answer: ", answers)
                                    st.write("Explanation: ", explanations)

                                text_to_download.add_heading(str(question_num) + ": " + scenario + questions, level=1)
                                text_to_download.add_paragraph(choices)
                                text_to_download.add_paragraph("Answer: " + answers)
                                text_to_download.add_paragraph("Explanation: " + explanations)
                                
                                question_num += 1

                            else: 

                                questions, choices, answers, explanations = prompts[question_type](format_instructions, chunks[i])
                                
                                st.write(str(question_num) + ": " + questions)
                                st.write(choices)
                                st.write("")
                                # st.write(answer)
                                with st.expander('Show me the answer 👀'): 
                                    st.write("Answer: ", answers)
                                    st.write("Explanation: ", explanations)

                                text_to_download.add_heading(str(question_num) + ": " + questions, level=1)
                                text_to_download.add_paragraph(choices)
                                text_to_download.add_paragraph("Answer: " + answers)
                                text_to_download.add_paragraph("Explanation: " + explanations)
                                # st.write(similar_docs[i])
                                
                                question_num += 1
                            continue
                        
                        except:
                            pass

                    doc_download = text_to_download
                    bio = io.BytesIO()
                    doc_download.save(bio)
                    if doc_download:
                        st.download_button(
                            label="Click here to download",
                            data=bio.getvalue(),
                            file_name="Quiz.docx",
                            mime="docx"
                        )
                    with st.sidebar.expander("token usage"):
                        st.write(cb)

    elif st.session_state["OpenAI_API_Key"] != "" and options == "Random questions":

        if pdf_docs != None:
            # with st.spinner("Wait for it..."):

            #Extract the text from the uploaded pdf file
            text = pdf_to_text(pdf_docs)

            #Split into chunks
            chunks = split_text(text, 750, 50)

            question_count = st.slider("Number of Questions", 1, len(chunks))

            st.write("Proportions must add up to 1")
            long = st.text_input("Proportion of scenario questions")
            short = st.text_input("Proportion of short questions")
            blank = st.text_input("Proportion of fill in blank questions")

            submit = st.button("Generate Quiz!")

            if submit:
                with get_openai_callback() as cb:
                        
                    text_to_download = Document()
                    indexes = random.sample(range(len(chunks)), question_count)

                    question_num = 1

                    for i in indexes:
                        try:
                            # question_type = random.randint(0,2)
                            question_type = np.random.choice(np.arange(0,3), p=[float(long), float(short), float(blank)])
                            # st.write(question_type)
                            format_instructions = formats[question_type]

                            if question_type == 0:

                                scenario, questions, choices, answers, explanations = prompts[question_type](format_instructions, chunks[i])
                                    
                                st.write(str(question_num) + ": " + scenario + " " + questions)
                                st.write(choices)
                                st.write("")
                                # st.write(answer)
                                with st.expander('Show me the answer 👀'): 
                                    st.write("Answer: ", answers)
                                    st.write("Explanation: ", explanations)

                                text_to_download.add_heading(str(question_num) + ": " + scenario + questions, level=1)
                                text_to_download.add_paragraph(choices)
                                text_to_download.add_paragraph("Answer: " + answers)
                                text_to_download.add_paragraph("Explanation: " + explanations)
                                
                                question_num += 1

                            else: 

                                questions, choices, answers, explanations = prompts[question_type](format_instructions, chunks[i])
                                
                                st.write(str(question_num) + ": " + questions)
                                st.write(choices)
                                st.write("")
                                # st.write(answer)
                                with st.expander('Show me the answer 👀'): 
                                    st.write("Answer: ", answers)
                                    st.write("Explanation: ", explanations)

                                text_to_download.add_heading(str(question_num) + ": " + questions, level=1)
                                text_to_download.add_paragraph(choices)
                                text_to_download.add_paragraph("Answer: " + answers)
                                text_to_download.add_paragraph("Explanation: " + explanations)
                                # st.write(similar_docs[i])
                                
                                question_num += 1

                            continue

                        except:
                            pass

                    doc_download = text_to_download
                    bio = io.BytesIO()
                    doc_download.save(bio)
                    if doc_download:
                        st.download_button(
                            label="Click here to download",
                            data=bio.getvalue(),
                            file_name="Quiz.docx",
                            mime="docx"
                        )
                    with st.sidebar.expander("token usage"):
                        st.write(cb)

    elif st.session_state["OpenAI_API_Key"] != "" and options == "Standard/Benchmark":

        if pdf_docs != None:
            # with st.spinner("Wait for it..."):

            #Extract the text from the uploaded pdf file
            text = pdf_to_text(pdf_docs)

            #Split into chunks
            chunks = split_text(text, 750, 50)

            question_count = st.slider("Number of Questions", 1, 10)

            quiz_description = st.text_area("Paste your learning goals here (i.e. benchmarks or standards) or just say 'give me some questions!'", key = "1")

            long = st.text_input("Proportion of scenario questions")
            short = st.text_input("Proportion of short questions")
            blank = st.text_input("Proportion of fill in blank questions")

            submit = st.button("Generate Quiz!")

            #Create embeddings
            embeddings = OpenAIEmbeddings(model_name="ada", openai_api_key=st.session_state["OpenAI_API_Key"])
            # embeddings = SentenceTransformerEmbeddings(model_name="all-MiniLM-L6-v2")

            #Store embeddings in a vectorstore
            index = Chroma.from_texts(chunks, embeddings)

            if submit:
                with get_openai_callback() as cb:
                        
                    llm = OpenAI(openai_api_key=st.session_state["OpenAI_API_Key"])
                    # llm=HuggingFaceHub(repo_id="bigscience/bloom", model_kwargs={"temperature":1e-10}, huggingfacehub_api_token="")

                    chain = load_qa_chain(llm, chain_type='stuff', verbose=False)
                    # qa = RetrievalQA.from_chain_type(llm=llm, chain_type="stuff", 
                    #                                  retriever=index.as_retriever(
                    #                                      search_kwargs={'k': question_count}
                    #                                      ),
                    #                                 )

                    similar_docs = index.similarity_search(quiz_description, k=question_count)
                    
                    text_to_download = Document()
                    question_num = 1

                    for i in range(question_count):

                        try:
                            question_type = np.random.choice(np.arange(0,3), p=[float(long), float(short), float(blank)])
                            # st.write(question_type)
                            # answer = qa.run(quiz_description)
                            answer = chain.run(input_documents = similar_docs, question=quiz_description)
                                
                            format_instructions = formats[question_type]
                            if question_type == 0:

                                scenario, questions, choices, answers, explanations = prompts[question_type](format_instructions, answer)
                                    
                                st.write(str(question_num) + ": " + scenario + " " + questions)
                                st.write(choices)
                                st.write("")
                                # st.write(answer)
                                with st.expander('Show me the answer 👀'): 
                                    st.write("Answer: ", answers)
                                    st.write("Explanation: ", explanations)

                                text_to_download.add_heading(str(question_num) + ": " + scenario + questions, level=1)
                                text_to_download.add_paragraph(choices)
                                text_to_download.add_paragraph("Answer: " + answers)
                                text_to_download.add_paragraph("Explanation: " + explanations)
                                
                                question_num += 1

                            else: 

                                questions, choices, answers, explanations = prompts[question_type](format_instructions, answer)
                                
                                st.write(str(question_num) + ": " + questions)
                                st.write(choices)
                                st.write("")
                                # st.write(answer)
                                with st.expander('Show me the answer 👀'): 
                                    st.write("Answer: ", answers)
                                    st.write("Explanation: ", explanations)

                                text_to_download.add_heading(str(question_num) + ": " + questions, level=1)
                                text_to_download.add_paragraph(choices)
                                text_to_download.add_paragraph("Answer: " + answers)
                                text_to_download.add_paragraph("Explanation: " + explanations)
                                # st.write(similar_docs[i])
                                
                                question_num += 1

                            continue

                        except:
                            pass

                    doc_download = text_to_download
                    bio = io.BytesIO()
                    doc_download.save(bio)
                    if doc_download:
                        st.download_button(
                            label="Click here to download",
                            data=bio.getvalue(),
                            file_name="Quiz.docx",
                            mime="docx"
                        )
                    with st.sidebar.expander("token usage"):
                        st.write(cb)

    elif st.session_state["OpenAI_API_Key"] != "" and options == "Key Terms":

        if pdf_docs != None:
            # with st.spinner("Wait for it..."):

            #Extract the text from the uploaded pdf file
            text = pdf_to_text(pdf_docs)

            #Split into chunks
            chunks = split_text(text, 500, 50)

            question_count = st.slider("Number of questions to generate for each term", 1, 3)

            key_terms = st.text_area("Paste your terms here. Each term should be seperated by a comma.", key = "1")

            long = st.text_input("Proportion of scenario questions")
            short = st.text_input("Proportion of short questions")
            blank = st.text_input("Proportion of fill in blank questions")

            submit = st.button("Generate Quiz!")

            #Create embeddings
            embeddings = OpenAIEmbeddings(model_name="ada", openai_api_key=st.session_state["OpenAI_API_Key"])
            # embeddings = SentenceTransformerEmbeddings(model_name="all-MiniLM-L6-v2")

            #Store embeddings in a vectorstore
            index = Chroma.from_texts(chunks, embeddings)

            if submit:
                with get_openai_callback() as cb:
                        
                    llm = OpenAI(openai_api_key=st.session_state["OpenAI_API_Key"])
                    # llm=HuggingFaceHub(repo_id="bigscience/bloom", model_kwargs={"temperature":1e-10}, huggingfacehub_api_token="")

                    chain = load_qa_chain(llm, chain_type='stuff', verbose=False)
                    # qa = RetrievalQA.from_chain_type(llm=llm, chain_type="stuff", 
                    #                                  retriever=index.as_retriever(
                    #                                      search_kwargs={'k': question_count}
                    #                                      ),
                    #                                 )
                    term_list = key_terms.split(",")

                    text_to_download = Document()     

                    question_num = 1

                    for term in term_list:


                        similar_docs = index.similarity_search(term, k=question_count)
                        

                        # indexes = random.sample(range(len(chunks)), question_count)
                        # iteration = 1
                        for i in range(question_count):

                            try:
                                question_type = np.random.choice(np.arange(0,3), p=[float(long), float(short), float(blank)])
                                # answer = qa.run(quiz_description)
                                answer = chain.run(input_documents = similar_docs, question=term)
                                    
                                format_instructions = formats[question_type]
                                if question_type == 0:

                                    scenario, questions, choices, answers, explanations = prompts[question_type](format_instructions, answer)
                                        
                                    st.write(str(question_num) + ": " + scenario + " " + questions)
                                    st.write(choices)
                                    st.write("")
                                    # st.write(answer)
                                    with st.expander('Show me the answer 👀'): 
                                        st.write("Answer: ", answers)
                                        st.write("Explanation: ", explanations)

                                    text_to_download.add_heading(str(question_num) + ": " + scenario + questions, level=1)
                                    text_to_download.add_paragraph(choices)
                                    text_to_download.add_paragraph("Answer: " + answers)
                                    text_to_download.add_paragraph("Explanation: " + explanations)
                                    
                                    question_num += 1

                                else: 

                                    questions, choices, answers, explanations = prompts[question_type](format_instructions, answer)
                                    
                                    st.write(str(question_num) + ": " + questions)
                                    st.write(choices)
                                    st.write("")
                                    # st.write(answer)
                                    with st.expander('Show me the answer 👀'): 
                                        st.write("Answer: ", answers)
                                        st.write("Explanation: ", explanations)

                                    text_to_download.add_heading(str(question_num) + ": " + questions, level=1)
                                    text_to_download.add_paragraph(choices)
                                    text_to_download.add_paragraph("Answer: " + answers)
                                    text_to_download.add_paragraph("Explanation: " + explanations)
                                    # st.write(similar_docs[i])
                                    
                                    question_num += 1

                                continue
                            except:
                                pass

                    doc_download = text_to_download
                    bio = io.BytesIO()
                    doc_download.save(bio)
                    if doc_download:
                        st.download_button(
                            label="Click here to download",
                            data=bio.getvalue(),
                            file_name="Quiz.docx",
                            mime="docx"
                        )
                    with st.sidebar.expander("token usage"):
                        st.write(cb)

    elif st.session_state["OpenAI_API_Key"] != "" and options == "Keyterm Finder":

        if pdf_docs != None:
            # with st.spinner("Wait for it..."):

            #Extract the text from the uploaded pdf file
            text = pdf_to_text(pdf_docs)

            #Split into chunks
            chunks = split_text(text, 3500, 50)

            num_keywords = st.slider("How many keywords do you want to extract?", 1, 100)
            
            num_docs = len(chunks)

            terms_per_chunk = round(num_keywords/num_docs)

            submit = st.button("Find Keywords!")

            keywords = ""

            if submit:
                with st.spinner("Wait for it..."):
                    with get_openai_callback() as cb:
                        i = 1
                        placeholder = st.empty()
                        for chunk in chunks:
                            placeholder.empty()
                            placeholder.text(str(i) + "/" + str(len(chunks)))
                            
                            format_instructions = format_schema_keywords(terms_per_chunk)
                            words, context = keyword_extractor(format_instructions, chunk, terms_per_chunk)

                            if len(keywords) == 0:
                                keywords += words
                            else:
                                keywords += ", " + words
                            
                            i += 1
                            
                            # st.write(context)
                            # kw_list = words.split(",")
                            # keywords.append(kw_list)
                    st.write(keywords)

                    # st.write(keywords)
                    with st.sidebar.expander("token usage"):
                        st.write(cb)
            

if __name__ == '__main__':
    main()